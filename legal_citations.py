# legal_citations.py
"""
Legal Citations System - Python Backend

For use in:
- Businessplan DOCX generation
- API validation
- Compliance checks

Mirrors TypeScript version for consistency
"""

from typing import Dict, List, Optional
from dataclasses import dataclass
from enum import Enum


class CitationCategory(str, Enum):
    """Citation categories"""
    ELIGIBILITY = "eligibility"
    AMOUNT = "amount"
    DURATION = "duration"
    HAUPTBERUFLICH = "hauptberuflich"
    NEBENTAETIGKEIT = "nebentaetigkeit"
    QUALIFICATIONS = "qualifications"
    WIRTSCHAFTLICH = "wirtschaftlich"
    DOCUMENTATION = "documentation"
    REPORTING = "reporting"


@dataclass
class LegalCitation:
    """Legal citation structure"""
    id: str
    title: str
    law: str
    section: str
    paragraph: Optional[str] = None
    sentence: Optional[str] = None
    short_text: str = ""
    full_text: Optional[str] = None
    official_source: str = ""
    date_valid: str = "2024-01-01"
    category: CitationCategory = CitationCategory.ELIGIBILITY
    tags: List[str] = None
    
    def __post_init__(self):
        if self.tags is None:
            self.tags = []
    
    def format(self) -> str:
        """Format citation for display"""
        result = f"{self.law} {self.section}"
        if self.paragraph:
            result += f" {self.paragraph}"
        if self.sentence:
            result += f" {self.sentence}"
        return result
    
    def to_docx_text(self) -> str:
        """Format for DOCX footnote/reference"""
        return f"{self.format()}: {self.short_text}"
    
    def to_dict(self) -> Dict:
        """Convert to dict for JSON/API"""
        return {
            'id': self.id,
            'title': self.title,
            'citation': self.format(),
            'short_text': self.short_text,
            'full_text': self.full_text,
            'official_source': self.official_source,
            'category': self.category.value,
            'tags': self.tags
        }


# ============================================================================
# LEGAL CITATIONS DATABASE
# ============================================================================

LEGAL_CITATIONS: Dict[str, LegalCitation] = {
    
    'gz_basic': LegalCitation(
        id='gz_basic',
        title='Gründungszuschuss Grundlage',
        law='SGB III',
        section='§ 93',
        short_text='Gründungszuschuss zur Aufnahme einer selbständigen Tätigkeit',
        full_text='Arbeitnehmerinnen und Arbeitnehmer, die durch Aufnahme einer selbständigen, hauptberuflichen Tätigkeit die Arbeitslosigkeit beenden, können zur Sicherung des Lebensunterhalts und zur sozialen Sicherung in der Zeit nach der Existenzgründung einen Gründungszuschuss erhalten.',
        official_source='https://www.gesetze-im-internet.de/sgb_3/__93.html',
        category=CitationCategory.ELIGIBILITY,
        tags=['grundlage', 'anspruch', 'selbständigkeit']
    ),
    
    'gz_amount_phase1': LegalCitation(
        id='gz_amount_phase1',
        title='Höhe Gründungszuschuss Phase 1',
        law='SGB III',
        section='§ 94',
        paragraph='Abs. 1',
        short_text='Höhe des individuellen ALG I + 300 EUR für 6 Monate',
        full_text='Der Gründungszuschuss wird in Höhe des zuletzt bezogenen Arbeitslosengeldes zur Sicherung des Lebensunterhalts und in Höhe von monatlich 300 Euro zur sozialen Absicherung für die Dauer von sechs Monaten geleistet.',
        official_source='https://www.gesetze-im-internet.de/sgb_3/__94.html',
        category=CitationCategory.AMOUNT,
        tags=['höhe', 'phase1', 'betrag']
    ),
    
    'gz_hauptberuflich': LegalCitation(
        id='gz_hauptberuflich',
        title='Hauptberufliche Selbständigkeit',
        law='SGB III',
        section='§ 93',
        paragraph='Abs. 2',
        short_text='Mind. 15 Stunden/Woche erforderlich für Hauptberuflichkeit',
        full_text='Hauptberuflich ist eine selbständige Tätigkeit, wenn sie zeitlich und wirtschaftlich die hauptsächliche Erwerbsgrundlage bildet. Dies ist in der Regel der Fall bei einem zeitlichen Umfang von mindestens 15 Stunden wöchentlich.',
        official_source='https://www.gesetze-im-internet.de/sgb_3/__93.html',
        category=CitationCategory.HAUPTBERUFLICH,
        tags=['hauptberuflich', '15stunden', 'zeitumfang']
    ),
    
    'gz_nebentaetigkeit_allowed': LegalCitation(
        id='gz_nebentaetigkeit_allowed',
        title='Nebentätigkeit während Gründungszuschuss',
        law='SGB III',
        section='§ 421',
        paragraph='i.V.m. § 155',
        short_text='Nebentätigkeit < 15h/Woche ist erlaubt, aber meldepflichtig',
        full_text='Eine Nebentätigkeit während des Bezugs von Gründungszuschuss ist grundsätzlich zulässig, solange sie die Hauptberuflichkeit der selbständigen Tätigkeit nicht gefährdet. Dies ist bei einem zeitlichen Umfang von weniger als 15 Stunden wöchentlich der Fall. Die Nebentätigkeit ist der Agentur für Arbeit unverzüglich zu melden.',
        official_source='https://www.gesetze-im-internet.de/sgb_3/__155.html',
        category=CitationCategory.NEBENTAETIGKEIT,
        tags=['nebentätigkeit', 'teilzeit', 'meldepflicht']
    ),
    
    'gz_teilzeit_warning': LegalCitation(
        id='gz_teilzeit_warning',
        title='Risiko bei zu hoher Nebentätigkeit',
        law='Fachliche Weisungen BA',
        section='zu § 93 SGB III',
        short_text='Bei > 15h/Woche oder > 50% Einkommen droht Rückforderung!',
        full_text='Überschreitet eine Nebentätigkeit 15 Stunden wöchentlich oder generiert sie mehr als 50% des Gesamteinkommens, gefährdet dies die Hauptberuflichkeit der selbständigen Tätigkeit. Dies kann zur Rückforderung des Gründungszuschusses führen.',
        official_source='https://www.arbeitsagentur.de/datei/fw-sgb-iii-93_ba014875.pdf',
        category=CitationCategory.NEBENTAETIGKEIT,
        tags=['warnung', 'rückforderung', 'risiko']
    ),
    
    'gz_tragfaehigkeit': LegalCitation(
        id='gz_tragfaehigkeit',
        title='Wirtschaftliche Tragfähigkeit',
        law='Fachliche Weisungen BA',
        section='zu § 93 SGB III',
        short_text='Geschäftsidee muss wirtschaftlich tragfähig sein (Finanzplan!)',
        full_text='Die geplante selbständige Tätigkeit muss eine tragfähige Existenzgrundlage bieten. Dies wird in der Regel anhand eines Businessplans mit Finanzplan geprüft. Der Finanzplan sollte zeigen, dass die Lebenshaltungskosten und Geschäftskosten mittelfristig aus dem Geschäft gedeckt werden können.',
        official_source='https://www.arbeitsagentur.de/datei/fw-sgb-iii-93_ba014875.pdf',
        category=CitationCategory.WIRTSCHAFTLICH,
        tags=['tragfähigkeit', 'finanzplan', 'businessplan']
    ),
    
    'gz_ermessen': LegalCitation(
        id='gz_ermessen',
        title='Gründungszuschuss als Ermessensleistung',
        law='SGB III',
        section='§ 93',
        paragraph='Abs. 1',
        short_text='KEIN Rechtsanspruch - Agentur entscheidet nach Ermessen',
        full_text='Der Gründungszuschuss ist eine Ermessensleistung. Die Arbeitnehmerin oder der Arbeitnehmer kann zur Aufnahme einer selbständigen, hauptberuflichen Tätigkeit einen Gründungszuschuss erhalten. Ein Rechtsanspruch besteht nicht.',
        official_source='https://www.gesetze-im-internet.de/sgb_3/__93.html',
        category=CitationCategory.ELIGIBILITY,
        tags=['ermessen', 'kein_anspruch', 'kann-leistung']
    ),
}


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def get_citation(citation_id: str) -> Optional[LegalCitation]:
    """Get citation by ID"""
    return LEGAL_CITATIONS.get(citation_id)


def get_citations_by_category(category: CitationCategory) -> List[LegalCitation]:
    """Get all citations by category"""
    return [c for c in LEGAL_CITATIONS.values() if c.category == category]


def get_citations_by_tag(tag: str) -> List[LegalCitation]:
    """Get all citations by tag"""
    return [c for c in LEGAL_CITATIONS.values() if tag in c.tags]


def format_citation_for_docx(citation_id: str) -> str:
    """
    Format citation for DOCX footnote
    
    Example:
    >>> format_citation_for_docx('gz_hauptberuflich')
    'SGB III § 93 Abs. 2: Mind. 15 Stunden/Woche erforderlich für Hauptberuflichkeit'
    """
    citation = get_citation(citation_id)
    if not citation:
        return ''
    return citation.to_docx_text()


def add_citation_to_text(text: str, citation_id: str) -> tuple[str, str]:
    """
    Add citation reference to text
    
    Returns:
        tuple: (text_with_reference, citation_text)
    
    Example:
    >>> text, citation = add_citation_to_text(
    ...     "Sie müssen mindestens 15 Stunden arbeiten",
    ...     "gz_hauptberuflich"
    ... )
    >>> print(text)
    'Sie müssen mindestens 15 Stunden arbeiten [1]'
    >>> print(citation)
    '[1] SGB III § 93 Abs. 2: Mind. 15 Stunden/Woche...'
    """
    citation = get_citation(citation_id)
    if not citation:
        return text, ''
    
    citation_text = f"[{citation.id}] {citation.to_docx_text()}"
    text_with_ref = f"{text} [{citation.format()}]"
    
    return text_with_ref, citation_text


# ============================================================================
# DOCX INTEGRATION
# ============================================================================

def add_citations_to_docx(doc, citations_used: List[str]):
    """
    Add citations section to DOCX document
    
    Usage in businessplan generator:
    >>> from docx import Document
    >>> doc = Document()
    >>> # ... generate businessplan ...
    >>> citations_used = ['gz_hauptberuflich', 'gz_nebentaetigkeit_allowed']
    >>> add_citations_to_docx(doc, citations_used)
    """
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    
    # Add section break
    doc.add_page_break()
    
    # Add heading
    heading = doc.add_heading('Rechtsgrundlagen', level=1)
    
    # Add each citation
    for citation_id in citations_used:
        citation = get_citation(citation_id)
        if not citation:
            continue
        
        # Citation title
        p = doc.add_paragraph()
        run = p.add_run(f"{citation.format()}: {citation.title}")
        run.bold = True
        run.font.size = Pt(11)
        
        # Citation text
        p = doc.add_paragraph(citation.full_text or citation.short_text)
        p.paragraph_format.left_indent = Pt(20)
        
        # Official source
        p = doc.add_paragraph()
        run = p.add_run('Quelle: ')
        run.italic = True
        run.font.size = Pt(9)
        run = p.add_run(citation.official_source)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = (102, 126, 234)  # Blue
        p.paragraph_format.left_indent = Pt(20)
        
        # Add spacing
        doc.add_paragraph()


# ============================================================================
# VALIDATION HELPERS
# ============================================================================

def validate_hauptberuflich(hours_per_week: int) -> tuple[bool, Optional[str]]:
    """
    Validate if hours meet Hauptberuflichkeit requirement
    
    Returns:
        tuple: (is_valid, citation_id if invalid)
    """
    if hours_per_week < 15:
        return False, 'gz_hauptberuflich'
    return True, None


def validate_nebentaetigkeit(
    part_time_hours: int,
    part_time_income: float,
    main_income: float
) -> tuple[bool, Optional[str]]:
    """
    Validate if part-time activity is allowed
    
    Returns:
        tuple: (is_valid, citation_id if invalid)
    """
    # Check hours
    if part_time_hours >= 15:
        return False, 'gz_teilzeit_warning'
    
    # Check income ratio
    total_income = part_time_income + main_income
    if total_income > 0 and (part_time_income / total_income) > 0.5:
        return False, 'gz_teilzeit_warning'
    
    return True, None


# ============================================================================
# EXAMPLE USAGE
# ============================================================================

if __name__ == '__main__':
    # Example: Get citation
    citation = get_citation('gz_hauptberuflich')
    print(f"Citation: {citation.format()}")
    print(f"Text: {citation.short_text}")
    print()
    
    # Example: Format for DOCX
    docx_text = format_citation_for_docx('gz_hauptberuflich')
    print(f"DOCX: {docx_text}")
    print()
    
    # Example: Validate
    is_valid, citation_id = validate_hauptberuflich(12)
    if not is_valid:
        citation = get_citation(citation_id)
        print(f"❌ Invalid: {citation.short_text}")
        print(f"   {citation.official_source}")
    
    # Example: Get all by category
    hauptberuflich_citations = get_citations_by_category(CitationCategory.HAUPTBERUFLICH)
    print(f"\nHauptberuflich citations: {len(hauptberuflich_citations)}")
    for c in hauptberuflich_citations:
        print(f"  - {c.format()}: {c.title}")
